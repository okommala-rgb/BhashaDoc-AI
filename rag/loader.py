"""
loader.py - Document loading module for BhashaDoc AI

Loads PDF, DOCX, and TXT files (from Streamlit UploadedFile objects or
filesystem paths) into LangChain Document objects, preserving metadata
such as filename and page number for every page/section.
"""

import io
import os
from typing import List, Union

from langchain_core.documents import Document
from pypdf import PdfReader
import docx


def _load_pdf(file_bytes: bytes, filename: str) -> List[Document]:
    """Load a PDF file, producing one Document per page with page metadata."""
    documents: List[Document] = []
    reader = PdfReader(io.BytesIO(file_bytes))
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        text = text.strip()
        if not text:
            continue
        documents.append(
            Document(
                page_content=text,
                metadata={
                    "filename": filename,
                    "page_number": page_number,
                    "source_type": "pdf",
                },
            )
        )
    return documents


def _load_docx(file_bytes: bytes, filename: str) -> List[Document]:
    """Load a DOCX file. DOCX has no native page concept, so content is
    split into pseudo-pages of fixed paragraph counts to keep citations
    meaningful for very large documents."""
    documents: List[Document] = []
    doc = docx.Document(io.BytesIO(file_bytes))

    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]

    # Include table content as well, since policy/legal docs often use tables.
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text)
            if row_text.strip():
                paragraphs.append(row_text)

    if not paragraphs:
        return documents

    paragraphs_per_page = 25  # pseudo-page boundary
    page_number = 1
    for i in range(0, len(paragraphs), paragraphs_per_page):
        chunk_paragraphs = paragraphs[i : i + paragraphs_per_page]
        text = "\n".join(chunk_paragraphs).strip()
        if not text:
            continue
        documents.append(
            Document(
                page_content=text,
                metadata={
                    "filename": filename,
                    "page_number": page_number,
                    "source_type": "docx",
                },
            )
        )
        page_number += 1

    return documents


def _load_txt(file_bytes: bytes, filename: str) -> List[Document]:
    """Load a plain text file. Splits into pseudo-pages by character count
    so very large text files still get sensible page-like metadata."""
    documents: List[Document] = []
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("utf-8", errors="ignore")

    chars_per_page = 3000
    page_number = 1
    for i in range(0, len(text), chars_per_page):
        page_text = text[i : i + chars_per_page].strip()
        if not page_text:
            continue
        documents.append(
            Document(
                page_content=page_text,
                metadata={
                    "filename": filename,
                    "page_number": page_number,
                    "source_type": "txt",
                },
            )
        )
        page_number += 1

    return documents


def _clean_text(text: str) -> str:
    """Basic cleaning pass: collapse whitespace, strip stray control chars."""
    text = text.replace("\x00", "")
    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]
    return "\n".join(lines)


def load_single_file(uploaded_file: Union[object, str]) -> List[Document]:
    """
    Load a single file into LangChain Documents.

    Accepts either:
      - a Streamlit UploadedFile-like object (has .name and .getvalue()/.read())
      - a filesystem path string
    """
    if isinstance(uploaded_file, str):
        filename = os.path.basename(uploaded_file)
        with open(uploaded_file, "rb") as f:
            file_bytes = f.read()
    else:
        filename = getattr(uploaded_file, "name", "unknown")
        if hasattr(uploaded_file, "getvalue"):
            file_bytes = uploaded_file.getvalue()
        else:
            file_bytes = uploaded_file.read()

    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == "pdf":
        docs = _load_pdf(file_bytes, filename)
    elif ext == "docx":
        docs = _load_docx(file_bytes, filename)
    elif ext == "txt":
        docs = _load_txt(file_bytes, filename)
    else:
        raise ValueError(f"Unsupported file type: {filename}")

    for d in docs:
        d.page_content = _clean_text(d.page_content)

    return [d for d in docs if d.page_content]


def load_documents(uploaded_files: List[Union[object, str]]) -> List[Document]:
    """Load multiple files (PDF/DOCX/TXT) into a single flat list of Documents."""
    all_documents: List[Document] = []
    for f in uploaded_files:
        all_documents.extend(load_single_file(f))
    return all_documents
